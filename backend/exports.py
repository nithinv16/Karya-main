"""Shared PDF / DOCX / XLSX export helpers for Karya reports.

Each generator returns raw bytes. The report is described as a list of "sections"
(heading + optional paragraphs / bullets / tables / images), which keeps the
per-resource endpoints tiny.
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Any, List, Optional

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
)

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side


BRAND = colors.HexColor("#EA580C")
INK = colors.HexColor("#09090B")
MUTED = colors.HexColor("#71717A")


@dataclass
class Section:
    heading: Optional[str] = None
    paragraphs: List[str] = field(default_factory=list)
    bullets: List[str] = field(default_factory=list)
    # `table` = [[row1_col1, row1_col2,...], ...] where first row is treated as header
    table: Optional[List[List[Any]]] = None
    # images: list of bytes (jpeg/png). Each rendered at ~5cm wide.
    images: List[bytes] = field(default_factory=list)


# ------------------------------------------------------------------- PDF

def build_pdf(title: str, subtitle: str, sections: List[Section]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.6 * cm, rightMargin=1.6 * cm,
        topMargin=1.6 * cm, bottomMargin=1.6 * cm,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName="Helvetica-Bold",
                        fontSize=20, textColor=INK, leading=24, spaceAfter=2)
    sub = ParagraphStyle("sub", parent=styles["Normal"], fontName="Helvetica",
                         fontSize=9, textColor=MUTED, leading=12, spaceAfter=16)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName="Helvetica-Bold",
                        fontSize=11, textColor=BRAND, leading=14, spaceBefore=10, spaceAfter=6,
                        textTransform="uppercase")
    body = ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica",
                          fontSize=10, textColor=INK, leading=14, spaceAfter=4)
    bullet = ParagraphStyle("bullet", parent=body, leftIndent=14, bulletIndent=4)

    flow = [Paragraph(title, h1), Paragraph(subtitle, sub)]
    for s in sections:
        if s.heading:
            flow.append(Paragraph(s.heading, h2))
        for p in s.paragraphs:
            flow.append(Paragraph(p, body))
        for b in s.bullets:
            flow.append(Paragraph(f"• {b}", bullet))
        if s.table:
            data = s.table
            tbl = Table(data, hAlign="LEFT", repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), INK),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 7),
                ("TOPPADDING", (0, 0), (-1, 0), 7),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAFAFA")]),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E4E4E7")),
            ]))
            flow.append(Spacer(1, 4))
            flow.append(tbl)
            flow.append(Spacer(1, 6))
        for img_bytes in s.images:
            try:
                flow.append(Spacer(1, 4))
                flow.append(Image(io.BytesIO(img_bytes), width=8 * cm, height=6 * cm, kind="proportional"))
            except Exception:
                pass
    doc.build(flow)
    return buf.getvalue()


# ------------------------------------------------------------------- DOCX

def build_docx(title: str, subtitle: str, sections: List[Section]) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    t = doc.add_paragraph()
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x09, 0x09, 0x0B)

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run(subtitle)
    sub_r.font.size = Pt(9)
    sub_r.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    for s in sections:
        if s.heading:
            h = doc.add_paragraph()
            hr = h.add_run(s.heading.upper())
            hr.bold = True
            hr.font.size = Pt(11)
            hr.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)
        for p in s.paragraphs:
            doc.add_paragraph(p)
        for b in s.bullets:
            doc.add_paragraph(b, style="List Bullet")
        if s.table:
            data = s.table
            tbl = doc.add_table(rows=len(data), cols=len(data[0]))
            tbl.style = "Light Grid Accent 1"
            for ri, row in enumerate(data):
                for ci, cell in enumerate(row):
                    tbl.rows[ri].cells[ci].text = str(cell)
                    if ri == 0:
                        for run in tbl.rows[ri].cells[ci].paragraphs[0].runs:
                            run.bold = True
        for img_bytes in s.images:
            try:
                doc.add_picture(io.BytesIO(img_bytes), width=Cm(8))
            except Exception:
                pass
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ------------------------------------------------------------------- XLSX

def build_xlsx(title: str, sheets: List[dict]) -> bytes:
    """sheets = [{'name': 'Sheet1', 'rows': [[...], ...], 'header': True}]"""
    wb = Workbook()
    wb.remove(wb.active)

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="09090B")
    center = Alignment(horizontal="left", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="E4E4E7")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for s in sheets:
        ws = wb.create_sheet(s.get("name", "Sheet")[:31] or "Sheet")
        rows = s.get("rows", [])
        has_header = s.get("header", True)
        for ri, row in enumerate(rows, start=1):
            for ci, val in enumerate(row, start=1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.alignment = center
                cell.border = border
                if ri == 1 and has_header:
                    cell.font = header_font
                    cell.fill = header_fill
        # auto-size columns (approximate)
        for col_idx, col in enumerate(ws.columns, start=1):
            max_len = 10
            for c in col:
                try:
                    v = str(c.value) if c.value is not None else ""
                    max_len = max(max_len, min(50, len(v) + 2))
                except Exception:
                    pass
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = max_len

    if not wb.sheetnames:
        ws = wb.create_sheet(title[:31] or "Sheet")
        ws["A1"] = "No data"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ------------------------------------------------------------------- helpers

MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def safe_filename(name: str, ext: str) -> str:
    keep = "".join(c if c.isalnum() or c in "-_ " else "_" for c in (name or "export")).strip()
    return f"{keep or 'export'}.{ext}"
